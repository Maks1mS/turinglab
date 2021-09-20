import pydot
import requests
import shutil
from os import path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from turinglab.emulator import Movemement, Program


def output(
    directory,
    program: Program,
    tests,
    docx_name="report",
    graph_name="graph"
):
    """Create output of a program and tests. It's includes a docx file with
    a 'command system', 'functional table' and tests.
    """

    def create_paragraph(document, text=""):
        p = document.add_paragraph(text)
        style_paragraph(p)
        return p

    def style_paragraph(p):
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def state_index(state):
        return str('z' if state == -1 else state)

    def add_state(p, state):
        p.add_run('q')
        index_text = p.add_run(state_index(state))
        index_text.font.subscript = True

    def get_symbol(x):
        return x if program.is_blank(x) is False else 'λ'

    document = Document()

    #
    # Command system
    #
    create_paragraph(document, "Система команд:")

    symbols = program.get_symbols()

    cols = len(symbols)
    rows = program.state_count
    program_table = [[[None] for n in range(cols)] for m in range(rows)]

    for i in range(rows):
        for j in range(cols):
            program_table[i][j] = program.get(i, symbols[j])

    for i in range(rows):
        for j in range(cols):
            action = program_table[i][j]
            if action is None:
                continue

            p = create_paragraph(document)
            add_state(p, i)

            movement = Movemement(action.movement).name

            p.add_run(get_symbol(symbols[j]) + ' → ')
            add_state(p, action.state)
            p.add_run(get_symbol(action.symbol) + movement)

    #
    # Functional table
    #
    create_paragraph(document)
    create_paragraph(document, "Функциональная таблица:")

    table = document.add_table(rows + 1, cols + 1)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True

    for i in range(0, rows):
        p = table.rows[i + 1].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p)
        add_state(p, i)

    for i, x in enumerate(symbols):
        p = table.rows[0].cells[i + 1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p)
        p.add_run(get_symbol(x))

    for i in range(rows):
        for j in range(cols):
            action = program_table[i][j]
            if action is None:
                continue

            p = table.rows[i + 1].cells[j + 1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p)

            movement = Movemement(action.movement).name
            add_state(p, action.state)
            p.add_run(get_symbol(action.symbol) + movement)

    #
    # Tests
    #
    p = create_paragraph(document)

    for i, test in enumerate(tests):
        p = create_paragraph(document)
        p.add_run(f'Тест {i + 1}')

        for j, data in enumerate(test):
            p = create_paragraph(document)
            p.add_run('K')

            index_text = p.add_run(str(j))
            index_text.font.subscript = True

            p.add_run(': ')

            head, state, tape = data

            tape_str = ''.join(tape.values())
            head -= tape.l_index

            left_str = tape_str[:head].lstrip(program.blank_symbol)
            left_str = left_str.replace(program.blank_symbol, 'λ')
            right_str = tape_str[head + 1:].rstrip(program.blank_symbol)
            right_str = right_str.replace(program.blank_symbol, 'λ')

            p.add_run(left_str)
            add_state(p, state)
            p.add_run(get_symbol(tape_str[head]))
            p.add_run(right_str)

    document.save(path.join(directory, docx_name + ".docx"))

    #
    # Image
    #

    g = pydot.Dot("mygraph", nodesep=0.75)

    g.set_node_defaults(
        width='1.5',
        fontsize='40',
        shape='circle'
    )
    g.set_edge_defaults(
        fontsize='20'
    )

    end_node = pydot.Node(f'gz', label='<g<SUB>z</SUB>>', shape='circle')
    g.add_node(end_node)

    for i in range(rows):
        node = pydot.Node(f'g{i}', label=f'<g<SUB>{i}</SUB>>', shape='circle')
        g.add_node(node)
        for j in range(cols):
            action = program_table[i][j]
            if action is None:
                continue

            movement = Movemement(action.movement).name

            label = get_symbol(symbols[j]) +\
                '/' + get_symbol(action.symbol) +\
                '/' + movement

            src = f'g{i}'
            dst = f'g{state_index(action.state)}'
            if src == dst:
                src = src + ':ne'
                dst = dst + ':se'

            g.add_edge(pydot.Edge(src, dst, label=label))

    dot_path = path.join(directory, f'{graph_name}.dot')
    g.write_raw(dot_path, encoding='utf-8')

    try:
        svg_path = path.join(directory, f'{graph_name}.svg')

        g.write_svg(svg_path, encoding='utf-8')
    except Exception:
        output_raw_dot = g.to_string()
        img_data = requests.get(
            "https://quickchart.io/graphviz?graph=" + output_raw_dot
        ).content

        svg_path = path.join(directory, f'{graph_name}_quickchart.svg')

        with open(svg_path, 'wb') as handler:
            handler.write(img_data)
